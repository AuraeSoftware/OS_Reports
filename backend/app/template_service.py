"""
Detects fillable fields in an uploaded .docx template and fills them back in.

Two placeholder styles are supported, and templates can mix both:
  1. Tag style   -> {{ Policy Number }}
  2. Label style -> "Policy Number:" followed by blank space / underscores,
                    either in a paragraph or as an empty table cell next to
                    a label cell (very common in insurance form templates).
"""
import re
from docx import Document

TAG_PATTERN = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}")
LABEL_PATTERN = re.compile(
    r"^(?P<label>[A-Za-z][A-Za-z0-9 /&\-\.\(\)]{1,60}?)\s*:\s*(?P<rest>_{2,}|\s*)$"
)


def slugify(text: str) -> str:
    key = re.sub(r"[^a-z0-9]+", "_", text.strip().lower()).strip("_")
    return key


def _humanize(key: str) -> str:
    return key.replace("_", " ").title()


def detect_fields(docx_path: str) -> list[dict]:
    """Scan the template and return the list of detected fields:
    [{"key": "policy_number", "label": "Policy Number", "style": "tag"}]
    """
    doc = Document(docx_path)
    fields: list[dict] = []
    seen: set[str] = set()

    def add(label: str, style: str):
        label = label.strip()
        key = slugify(label)
        if key and key not in seen:
            seen.add(key)
            fields.append({"key": key, "label": label, "style": style})

    def scan_tag_text(text: str):
        for m in TAG_PATTERN.finditer(text):
            add(m.group(1), "tag")

    def scan_label_paragraph(text: str):
        text = text.strip()
        if not text or "{{" in text:
            return
        m = LABEL_PATTERN.match(text)
        if m:
            add(m.group("label"), "label")

    # top-level paragraphs
    for p in doc.paragraphs:
        scan_tag_text(p.text)
        scan_label_paragraph(p.text)

    # tables: tag-style anywhere, label-style as "label cell" + "empty next cell"
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            for cell in cells:
                for p in cell.paragraphs:
                    scan_tag_text(p.text)
            for i in range(len(cells) - 1):
                label_text = cells[i].text.strip()
                next_text = cells[i + 1].text.strip()
                if label_text.endswith(":") and not next_text and len(label_text) > 1:
                    add(label_text.rstrip(":").strip(), "label")

    return fields


def _set_paragraph_text(p, new_text: str):
    """Clear all runs and put the full replacement text in the first one.
    Loses per-run formatting inside the paragraph but keeps paragraph style,
    which is an acceptable tradeoff for form field filling."""
    if p.runs:
        p.runs[0].text = new_text
        for run in p.runs[1:]:
            run.text = ""
    else:
        p.text = new_text


def fill_template(template_path: str, data: dict, output_path: str) -> str:
    """data is {field_key: value}. Any placeholder without a matching key
    (or with an empty value) is left blank rather than showing a stray tag."""
    doc = Document(template_path)

    def resolved(raw_label: str):
        key = slugify(raw_label)
        val = data.get(key, "")
        return "" if val is None else str(val)

    def fill_tag_paragraph(p):
        full_text = p.text
        if "{{" not in full_text:
            return
        new_text = TAG_PATTERN.sub(lambda m: resolved(m.group(1)), full_text)
        if new_text != full_text:
            _set_paragraph_text(p, new_text)

    def fill_label_paragraph(p):
        text = p.text.strip()
        if not text or "{{" in text:
            return
        m = LABEL_PATTERN.match(text)
        if not m:
            return
        label = m.group("label").strip()
        key = slugify(label)
        if key in data and data[key] not in (None, ""):
            _set_paragraph_text(p, f"{label}: {data[key]}")

    for p in doc.paragraphs:
        fill_tag_paragraph(p)
        fill_label_paragraph(p)

    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            for cell in cells:
                for p in cell.paragraphs:
                    fill_tag_paragraph(p)
            for i in range(len(cells) - 1):
                label_text = cells[i].text.strip()
                if label_text.endswith(":") and len(label_text) > 1:
                    label = label_text.rstrip(":").strip()
                    key = slugify(label)
                    target_cell = cells[i + 1]
                    if (
                        key in data
                        and data[key] not in (None, "")
                        and not target_cell.text.strip()
                        and target_cell.paragraphs
                    ):
                        _set_paragraph_text(target_cell.paragraphs[0], str(data[key]))

    doc.save(output_path)
    return output_path
